from __future__ import annotations

import csv
import json
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from app.core.logger import get_logger

logger = get_logger(__name__)


class ExportService:
    def __init__(self, output_dir: str | Path) -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export_json(self, rows: Iterable[dict[str, Any]], name: str) -> Path:
        path = self.output_dir / f"{name}.json"
        path.write_text(
            json.dumps(list(rows), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        return path

    def export_csv(self, rows: list[dict[str, Any]], name: str) -> Path:
        path = self.output_dir / f"{name}.csv"
        if not rows:
            path.write_text("", encoding="utf-8-sig")
            return path
        with path.open("w", encoding="utf-8-sig", newline="") as file:
            writer = csv.DictWriter(file, fieldnames=list(rows[0].keys()))
            writer.writeheader()
            writer.writerows(rows)
        return path

    def export_excel(self, rows: list[dict[str, Any]], name: str) -> Path:
        try:
            from openpyxl import Workbook
        except ImportError as exc:
            raise RuntimeError("openpyxl is required for Excel export") from exc
        path = self.output_dir / f"{name}.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = name[:31] or "Sheet"
        if rows:
            headers = list(rows[0].keys())
            sheet.append(headers)
            for row in rows:
                sheet.append([row.get(header) for header in headers])
        workbook.save(path)
        return path

    def export_docx(
        self, rows: list[dict[str, Any]], name: str, *, title: str = "Export"
    ) -> Path:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX export") from exc
        path = self.output_dir / f"{name}.docx"
        document = Document()
        document.add_heading(title, level=1)
        for row in rows:
            document.add_paragraph(" | ".join(f"{key}: {value}" for key, value in row.items()))
        document.save(path)
        return path

    def list_files(self) -> list[dict[str, Any]]:
        files = []
        for path in sorted(self.output_dir.iterdir(), key=lambda p: p.stat().st_mtime, reverse=True):
            if path.is_file():
                files.append(
                    {
                        "name": path.name,
                        "path": str(path),
                        "size": path.stat().st_size,
                        "modified": path.stat().st_mtime,
                    }
                )
        return files

    def resolve_path(self, name: str) -> Path:
        path = (self.output_dir / name).resolve()
        if not path.is_relative_to(self.output_dir.resolve()):
            raise ValueError("path escapes output directory")
        return path

    def delete_file(self, name: str) -> bool:
        path = self.resolve_path(name)
        if not path.exists():
            return False
        path.unlink()
        return True
